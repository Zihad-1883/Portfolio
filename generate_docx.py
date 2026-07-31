import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

doc = docx.Document()

# Set margins matching the original document
for section in doc.sections:
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

COLOR_BLACK = RGBColor(0, 0, 0)
COLOR_BLUE = RGBColor(0, 0, 255)

def set_para_spacing(p, before=0, after=1, line_spacing=1.05):
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line_spacing

def add_heading(text):
    p = doc.add_paragraph()
    set_para_spacing(p, 4, 1, 1.0)
    r = p.add_run(text.upper())
    r.font.name = 'Arial'
    r.font.size = Pt(9.5)
    r.font.bold = True
    r.font.color.rgb = COLOR_BLACK

def add_body(text):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 1.5, 1.05)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(8.5)
    r.font.color.rgb = COLOR_BLACK

def add_skill(label, items):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 1, 1.05)
    r1 = p.add_run(label + ": ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(8.5)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_BLACK
    r2 = p.add_run(items)
    r2.font.name = 'Arial'
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = COLOR_BLACK

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    set_para_spacing(p, 0, 1, 1.05)
    p.paragraph_format.left_indent = Inches(0.18)
    r = p.add_run(text)
    r.font.name = 'Arial'
    r.font.size = Pt(8.5)
    r.font.color.rgb = COLOR_BLACK

def add_item_title(left_title, date_str=""):
    p = doc.add_paragraph()
    set_para_spacing(p, 2.5, 0.5, 1.0)
    r1 = p.add_run(left_title)
    r1.font.name = 'Arial'
    r1.font.size = Pt(8.8)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_BLACK
    if date_str:
        r2 = p.add_run(" | " + date_str)
        r2.font.name = 'Arial'
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = COLOR_BLACK

def add_project_title(title, overview, live_url, github_url):
    p = doc.add_paragraph()
    set_para_spacing(p, 2.5, 0.5, 1.0)
    
    r1 = p.add_run(title + " – " + overview + " | ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(8.8)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_BLACK
    
    r2 = p.add_run("Live")
    r2.font.name = 'Arial'
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = COLOR_BLUE
    r2.font.underline = True
    
    r3 = p.add_run(" | ")
    r3.font.name = 'Arial'
    r3.font.size = Pt(8.5)
    r3.font.color.rgb = COLOR_BLACK
    
    r4 = p.add_run("GitHub")
    r4.font.name = 'Arial'
    r4.font.size = Pt(8.5)
    r4.font.color.rgb = COLOR_BLUE
    r4.font.underline = True

def add_tech_stack(stack_str):
    p = doc.add_paragraph()
    set_para_spacing(p, 0, 1, 1.0)
    r1 = p.add_run("Tech Stack: ")
    r1.font.name = 'Arial'
    r1.font.size = Pt(8.2)
    r1.font.bold = True
    r1.font.color.rgb = COLOR_BLACK
    
    r2 = p.add_run(stack_str)
    r2.font.name = 'Arial'
    r2.font.size = Pt(8.2)
    r2.font.italic = True
    r2.font.color.rgb = COLOR_BLACK

# HEADER
p_head = doc.add_paragraph()
p_head.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_para_spacing(p_head, 0, 1, 1.0)

r_name = p_head.add_run("MIZBAUR RAHMAN ZIHAD\n")
r_name.font.name = 'Arial'
r_name.font.size = Pt(14)
r_name.font.bold = True
r_name.font.color.rgb = COLOR_BLACK

r_title = p_head.add_run("Junior Full-Stack Developer\n")
r_title.font.name = 'Arial'
r_title.font.size = Pt(9.5)
r_title.font.bold = True
r_title.font.color.rgb = COLOR_BLACK

r_contact = p_head.add_run("Sylhet, Bangladesh | +8801865662482 | zihad.pilot18@gmail.com | LinkedIn | GitHub | Portfolio")
r_contact.font.name = 'Arial'
r_contact.font.size = Pt(8.2)
r_contact.font.color.rgb = COLOR_BLUE
r_contact.font.underline = True

# CAREER OBJECTIVE
add_heading("CAREER OBJECTIVE")
add_body("I am a passionate web developer aiming to grow into a full-stack and backend developer, continuously exploring modern technologies along the way. I aspire to contribute to impactful projects in the tech industry while expanding my skills in system design and scalable architecture. In the long term, I plan to venture into AI and Machine Learning, which aligns naturally with my academic background in Mathematics.")

# TECHNICAL SKILLS
add_heading("TECHNICAL SKILLS")
add_skill("Frontend", "HTML5, CSS3, JavaScript ES6+, TypeScript, React.js, Next.js, Tailwind CSS, DaisyUI, HeroUI, Framer Motion, responsive design")
add_skill("Backend", "Node.js, Express.js, PostgreSQL, Prisma ORM, MongoDB, NeonDB, REST API, CRUD operations, JWT, Better Auth, Bcrypt, Modular Architecture, RBAC")
add_skill("Tools / Deployment", "Git, GitHub, VS Code, Vite, npm, Postman, Figma, Vercel, Netlify, Render")
add_skill("AI Workflow", "ChatGPT, Gemini, Claude, Antigravity, AI-assisted debugging, documentation support, prompt engineering")
add_skill("Soft Skills", "Logical reasoning, Problem Solving, growth mindset")

# PRACTICAL EXPERIENCE
add_heading("PRACTICAL EXPERIENCE")
add_item_title("Back-End AI Engineering Intern — FlyRank AI", "July 2026 – Sept 2026")
add_bullet("Architected backend API routes using Express.js & Node.js to stream LLM responses and manage asynchronous background tasks.")
add_bullet("Integrated vector search workflows and AI prompt pipelines utilizing PostgreSQL, TypeScript, and cloud infrastructure.")
add_bullet("Enforced strict API contract testing, schema validation, and remote-first version control standards.")

add_item_title("Complete Web Development Course — Programming Hero", "Dec 2025 – Jun 2026")
add_bullet("Completed extensive project-based training covering full-stack JavaScript, React, Node.js, Express, MongoDB, secure deployment, and version-controlled Git workflows.")
add_bullet("Developed production-ready frontend and full-stack projects featuring protected router states and smooth user authentication flows.")

add_item_title("Next Level AI-Driven Software Engineering Bootcamp — Programming Hero", "Apr 2026 – Present")
add_bullet("Mastered TypeScript foundations, advanced types, data constraints, raw SQL connection pools, and serverless PostgreSQL via NeonDB & Prisma ORM.")

# TECHNICAL PROJECTS
add_heading("TECHNICAL PROJECTS")

add_project_title("GearUp", "Sports & Outdoor Equipment Rental API", "https://gearup-backend-4eca.onrender.com/", "https://github.com/Zihad-1883/L2-A4-Gear-Up")
add_tech_stack("Node.js, Express.js, TypeScript, PostgreSQL, Prisma ORM, SSLCommerz, JWT, Bcrypt")
add_bullet("Built a peer-to-peer equipment rental API supporting a 5-stage order lifecycle state machine (PENDING ➔ APPROVED ➔ PAID ➔ PICKED_UP ➔ RETURNED).")
add_bullet("Implemented backend booking collision-checking algorithms and multi-role RBAC separating CUSTOMER, PROVIDER, and ADMIN permissions.")
add_bullet("Integrated two-step SSLCommerz payment validation with server-to-server verification APIs and IPN webhooks.")

add_project_title("ErythroShare", "Healthcare & Blood Donation Platform", "https://l1-a10-erythro-share-client.vercel.app/", "https://github.com/Zihad-1883/L1-A10-ErythroShare-Client")
add_tech_stack("Next.js 16 (App Router), Express.js, MongoDB, Better Auth, Tailwind CSS, JWT")
add_bullet("Engineered a multi-role healthcare coordination portal managing permissions for ADMIN, VOLUNTEER, and DONOR users.")
add_bullet("Built protected Express REST API endpoints with session caching and real-time blood request dispatch tracking.")
add_bullet("Implemented donor search indexing by blood group, district, and availability status to streamline emergency requests.")

add_project_title("NextKey", "Property Rental Platform", "https://project-next-key.vercel.app/", "https://github.com/Zihad-1883/Project-Next-Key")
add_tech_stack("Next.js 16 (App Router), Express 5, TypeScript, MongoDB, Monorepo Architecture")
add_bullet("Designed a full-stack monorepo property rental platform sharing TypeScript interfaces between Next.js App Router client and Express TS backend.")
add_bullet("Developed a dynamic multi-parameter database query engine for property filtering (price range, category, bed/bath count, location).")
add_bullet("Constructed landlord rental management dashboards with request tracking and tenant approval workflows.")

# EDUCATION
add_heading("EDUCATION")
add_item_title("B.Sc. in Mathematics (3rd Year) – Shahjalal University of Science and Technology (SUST), Sylhet", "2023 – Present")
add_bullet("Core Focus: Analytical Problem Solving, Discrete Mathematics, Logic Optimization, Algorithms.")

# CERTIFICATIONS & LANGUAGES
add_heading("CERTIFICATIONS & LANGUAGES")
add_skill("Certifications", "Complete Web Development with MERN Stack – Programming Hero (2026, Pending); Next Level Software Engineering Certification (Ongoing).")
add_skill("Languages", "Bangla – Native; English – Intermediate Professional Working Proficiency, comfortable with technical documentation, code reviews, and developer communication.")

doc_path = r"e:\001_Zihad-Programming\Portfolio\public\RESUME_UPDATED.docx"
doc.save(doc_path)
doc.save(r"e:\001_Zihad-Programming\Portfolio\public\RESUME.docx")
print("Saved DOCX files matching reference image shape.")
